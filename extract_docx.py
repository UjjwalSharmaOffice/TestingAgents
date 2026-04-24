import sys
from docx import Document

files = [
    r'c:\Users\UjjwalSharma\Desktop\TestingAgents\epics\CodeVal Services epics and User stories.docx',
    r'c:\Users\UjjwalSharma\Desktop\TestingAgents\templates\PID_TestPlan_example.docx',
    r'c:\Users\UjjwalSharma\Desktop\TestingAgents\templates\Test Strategy_template (single document) (3).docx'
]

def extract_text(file_path):
    try:
        doc = Document(file_path)
        print(f'--- Content of: {file_path} ---')
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                print(' | '.join(row_text))
        print('\n' + '='*50 + '\n')
    except Exception as e:
        print(f'Error reading {file_path}: {e}')

for file in files:
    extract_text(file)
