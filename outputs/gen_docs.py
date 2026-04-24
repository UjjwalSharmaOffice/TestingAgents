"""
Generate CodEval Test Strategy and Test Plan .docx documents
based on the existing templates and qa_output.json.
"""
import json
import copy
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(r'c:\Users\UjjwalSharma\Desktop\TestingAgents')
TEMPLATES = BASE / 'templates'
OUTPUTS = BASE / 'outputs'

with open(OUTPUTS / 'qa_output.json', encoding='utf-8') as f:
    qa = json.load(f)

features = qa['features']

# ─────────────────────────────────────────────
# helpers
# ─────────────────────────────────────────────
def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, style=None):
    if style:
        p = doc.add_paragraph(text, style=style)
    else:
        p = doc.add_paragraph(text)
    return p

def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def set_col_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def add_table_row(table, cells_text, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells_text)):
        cell.text = text
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

# ─────────────────────────────────────────────────────────────────────────────
# TEST STRATEGY
# ─────────────────────────────────────────────────────────────────────────────
def build_test_strategy():
    src = TEMPLATES / 'Test Strategy_template (single document) (3).docx'
    dst = OUTPUTS / 'CodEval_Test_Strategy.docx'
    shutil.copy2(src, dst)
    doc = Document(dst)

    # ── clear body paragraphs after cover / TOC (keep first 3 paras = title area)
    # We rebuild from scratch after the TOC heading
    to_remove = []
    found_toc = False
    for p in doc.paragraphs:
        if found_toc:
            to_remove.append(p._element)
        if 'Contents' in p.text or p.style.name == 'TOC Heading':
            found_toc = True

    for el in to_remove:
        el.getparent().remove(el)

    # Also remove all tables (we will re-add what we need)
    for tbl in list(doc.tables):
        tbl._element.getparent().remove(tbl._element)

    # ── Section 1: Test Strategy ──────────────────────────────────────────────
    add_heading(doc, '1. Test Strategy', 1)
    add_para(doc, (
        'This document describes the test strategy for the CodEval Platform — an EPAM-internal '
        'coding-assessment system consisting of an Admin Platform (Team 1) and a Candidate Platform '
        '(Team 2). The scope of this strategy covers all Team 1 (Authentication + Admin Platform) '
        'features and the agreed integration contracts shared between Team 1 and Team 2. '
        'The goal is to validate functional correctness, API contract compliance, security of '
        'authentication flows, and end-to-end integration between both teams.'
    ))

    # ── Section 2: Test Strategy Outline ─────────────────────────────────────
    add_heading(doc, '2. Test Strategy Outline', 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    hdr[0].text = 'Section'
    hdr[1].text = 'Purpose'
    hdr[2].text = 'Notes'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    outline_rows = [
        ('Scope of testing',
         'Authentication Module (Login, Signup, Password Reset)\n'
         'Admin Platform: User & Batch Mgmt, Assessment Config & Invite, '
         'Question Mgmt + AI Generation, Email Notifications, Statistics\n'
         'Integration contracts: /assessments/{id}/config, /questions/{id}, '
         '/assessments/{id}/testcases, /results/* APIs, WebSocket RUN/OUTPUT/DONE',
         'Covers all E1, A-E1–A-E5, INT epics'),
        ('Out of scope',
         'Candidate Platform (Team 2) internal features (B-E1–B-E5) except where they form '
         'integration contract endpoints consumed by Team 1.',
         'Team 2 internal features tested separately'),
        ('Acceptance criteria for release',
         '- No open Critical or Major defects at release\n'
         '- All happy-path test scenarios pass\n'
         '- All integration contract API tests pass with correct payloads\n'
         '- RBAC validation confirmed on all protected endpoints',
         ''),
        ('Test types',
         'Functional, API, Integration, Smoke, Regression, UI',
         'See Section 3'),
        ('Test phases',
         'Phase 1: Smoke — verify core auth and admin shell\n'
         'Phase 2: Feature — all Team 1 stories\n'
         'Phase 3: Integration — cross-team API contracts\n'
         'Phase 4: Regression — re-run after bug fixes',
         ''),
        ('Test automation',
         'Smoke and API contract tests automated (REST Assured / Postman collections).\n'
         'UI functional tests manual; automation candidates identified post-Phase 2.',
         ''),
        ('Test environments',
         'DEV: local Docker Compose\n'
         'QA: shared staging environment with both team deployments\n'
         'PROD: not in scope for QA',
         ''),
        ('Test tools',
         'JIRA (bug tracking), Postman / REST Assured (API), Selenium/Playwright (UI smoke), '
         'QA Space (test documentation)',
         ''),
        ('Risks',
         '- Integration contract API schemas may drift between teams\n'
         '- AI generation (Spring AI) depends on LLM availability\n'
         '- HMAC token expiry timing sensitive to clock sync between services',
         'Mitigations defined in Risk section of Test Plan'),
    ]
    for row_data in outline_rows:
        row = tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ── Section 3: Testing Types ──────────────────────────────────────────────
    add_heading(doc, '3. Testing Types', 1)

    types = [
        ('3.1 Requirements Testing',
         'Verify that all user stories and acceptance criteria in the CodEval epics are fully '
         'implemented and traceable to test scenarios.',
         [('Test Objective', 'Eliminate gaps, contradictions, and ambiguity in requirements coverage.'),
          ('Key Considerations', 'Each feature in qa_output.json maps 1-to-1 with a QA story and acceptance criteria.')]),
        ('3.2 Feature Testing',
         'Verify that each feature of the CodEval Admin Platform and Authentication module operates '
         'in accordance with its acceptance criteria.',
         [('Test Objective', 'Validate happy path and all explicitly-stated negative scenarios per feature.'),
          ('Key Considerations', 'Test scenarios in qa_output.json are the primary test-case source.')]),
        ('3.3 Ad-hoc / Exploratory Testing',
         'Explore the system to discover defects not covered by scripted test cases, focusing on '
         'auth edge cases, admin workflow combinations, and AI generation variability.',
         [('Test Objective', 'Break the system; find unexpected behaviour in multi-step admin workflows.'),
          ('Key Considerations', 'Prioritise auth flows, invite + assessment lifecycle, and AI question edge states.')]),
        ('3.4 UI Testing',
         'Verify that the Admin Platform UI meets design expectations and that all controls, '
         'inputs, and navigations function as specified.',
         [('Test Objective', 'Validate UI layouts, Monaco editor integration, filter/pagination controls, modals.'),
          ('Key Considerations', 'Login page, User Management, Assessment Management, Question Editor pages.')]),
        ('3.5 Smoke Testing',
         'Rapidly assess build readiness by verifying the most critical paths: login, admin dashboard '
         'load, user list, and assessment creation.',
         [('Test Objective', 'No Critical defects block further testing.'),
          ('Key Considerations', 'Run on every new QA build; includes TS-E1-01, TS-A1-01, TS-A3-01.')]),
        ('3.6 Compatibility Testing',
         'Verify that the Admin Platform works correctly on the required browsers and screen sizes.',
         [('Test Objective', 'Consistent functionality across Chrome, Firefox, and Edge (latest two versions).'),
          ('Key Considerations', 'Monaco editor compatibility is a key risk area.')]),
        ('3.7 Regression Testing & Re-Testing',
         'Re-execute previously failed tests after fixes; confirm no regressions in unchanged areas.',
         [('Test Objective', 'Confirm fixes and detect regressions introduced by code changes.'),
          ('Key Considerations', 'Full regression suite run before each release candidate.')]),
        ('3.8 API Testing',
         'Verify that all REST API endpoints and WebSocket message contracts behave per specification.',
         [('Test Objective', 'Validate request/response payloads, status codes, RBAC enforcement, and error responses.'),
          ('Key Considerations', 'Integration contract endpoints (INT epics) are highest priority:\n'
           'GET /assessments/{id}/config, GET /questions/{id}, GET /assessments/{id}/testcases,\n'
           'GET /results/assessment/{id}, GET /results/user/{id}, WebSocket RUN/OUTPUT/DONE.')]),
        ('3.9 Integration Testing',
         'Verify interactions between Team 1 and Team 2 components across agreed API contracts.',
         [('Test Objective', 'End-to-end flow: admin creates assessment → invite sent → candidate session reads config → results returned to admin analytics.'),
          ('Key Considerations', 'Requires both Team 1 and Team 2 services running in QA environment.')]),
    ]

    for title, intro, table_rows in types:
        add_heading(doc, title, 2)
        add_para(doc, intro)
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.style = 'Table Grid'
        for label, value in table_rows:
            row = tbl2.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True

    # ── Section 4: Entry Criteria ─────────────────────────────────────────────
    add_heading(doc, '4. Entry Criteria', 1)
    add_para(doc, 'The Testing Team may suspend testing if any of the following occurs:')
    entry = [
        'It is impossible to install the new build successfully.',
        'There is a fault with a feature that prevents its testing.',
        'Application does not contain the specified change(s).',
        'New claimed functionality does not work or works improperly.',
        'A severe problem has occurred that does not allow testing to continue.',
        'Development has not corrected problems that previously suspended testing.',
        'A new version of the software is available to test.',
    ]
    for item in entry:
        add_bullet(doc, item)

    # ── Section 5: Bug and Documentation Tracking ─────────────────────────────
    add_heading(doc, '5. Bug and Documentation Tracking', 1)
    add_para(doc, (
        'Bug reporting is managed in Atlassian JIRA. Test documentation is tracked in QA Space '
        '(Jira plugin). Bug metrics and statistics are included in weekly test result reports.'
    ))
    add_heading(doc, '5.1 Bug Severity Definitions', 2)
    severities = [
        ('Critical', 'Application, component, or module crash or inaccessible.'),
        ('Major', 'Data corruption/loss, problem in major functionality, no workaround known.'),
        ('Medium', 'Problem with a workaround; secondary features do not work properly.'),
        ('Minor', 'Cosmetic flaw.'),
    ]
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.style = 'Table Grid'
    tbl3.rows[0].cells[0].text = 'Severity'
    tbl3.rows[0].cells[1].text = 'Definition'
    for cell in tbl3.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for sev, defn in severities:
        row = tbl3.add_row()
        row.cells[0].text = sev
        row.cells[1].text = defn

    doc.save(dst)
    print(f'Test Strategy saved: {dst}')

# ─────────────────────────────────────────────────────────────────────────────
# TEST PLAN
# ─────────────────────────────────────────────────────────────────────────────
def build_test_plan():
    src = TEMPLATES / 'PID_TestPlan_example.docx'
    dst = OUTPUTS / 'CodEval_Test_Plan.docx'
    shutil.copy2(src, dst)
    doc = Document(dst)

    # Remove body content after TOC, keep cover metadata
    to_remove = []
    found_toc = False
    for p in doc.paragraphs:
        if found_toc:
            to_remove.append(p._element)
        if 'Contents' in p.text or p.style.name == 'TOC Heading':
            found_toc = True
    for el in to_remove:
        el.getparent().remove(el)
    for tbl in list(doc.tables):
        tbl._element.getparent().remove(tbl._element)

    # ── 1. Introduction ───────────────────────────────────────────────────────
    add_heading(doc, '1. Introduction', 1)
    add_para(doc, (
        'This document describes the approach and methodologies used by the testing team to plan, '
        'organise, and perform testing of the CodEval Platform — an EPAM-internal coding-assessment '
        'system. The platform consists of an Authentication Module (shared), an Admin Platform '
        '(Team 1), and a Candidate Platform (Team 2). This test plan covers Team 1 features '
        'and the integration contracts shared between Team 1 and Team 2.'
    ))

    # ── 2. Scope of Work ──────────────────────────────────────────────────────
    add_heading(doc, '2. Scope of Work', 1)
    add_heading(doc, '2.1 Components and Functions to be Tested', 2)

    in_scope_tbl = doc.add_table(rows=1, cols=4)
    in_scope_tbl.style = 'Table Grid'
    hdr = in_scope_tbl.rows[0].cells
    for c, t in zip(hdr, ['#', 'Component / Module', 'Function', 'Reference']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True

    in_scope = [
        ('1', 'Authentication (Shared)', 'Login (EPAM @epam.com email + password, JWT, RBAC)', 'Epic E1 — TS-E1-01 to TS-E1-04'),
        ('2', 'Authentication (Shared)', 'Signup (EPAM email, password strength, role=USER)', 'Epic E1 — TS-E1-05 to TS-E1-07'),
        ('3', 'Authentication (Shared)', 'Password Reset (HMAC link, 15-min expiry, Redis session invalidation)', 'Epic E1 — TS-E1-08 to TS-E1-10'),
        ('4', 'Admin Platform — User & Batch Mgmt', 'Paginated/filterable user list (A-1.1)', 'Epic A-E1 — TS-A1-01 to TS-A1-03'),
        ('5', 'Admin Platform — User & Batch Mgmt', 'User detail + performance statistics (A-1.2)', 'Epic A-E1 — TS-A1-04'),
        ('6', 'Admin Platform — User & Batch Mgmt', 'Create / edit / deactivate users (A-1.3)', 'Epic A-E1 — TS-A1-05 to TS-A1-07'),
        ('7', 'Admin Platform — User & Batch Mgmt', 'Batch CRUD + CSV bulk-add + duplicate prevention (A-1.4)', 'Epic A-E1 — TS-A1-08 to TS-A1-10'),
        ('8', 'Admin Platform — Assessment Config', 'Create & configure assessment in DRAFT (A-3.1)', 'Epic A-E3 — TS-A3-01 to TS-A3-02'),
        ('9', 'Admin Platform — Assessment Config', 'Send invite emails; HMAC tokens; status=ACTIVE (A-3.2)', 'Epic A-E3 — TS-A3-03 to TS-A3-04'),
        ('10', 'Admin Platform — Assessment Config', 'View/edit/delete assessments; lifecycle guards (A-3.3)', 'Epic A-E3 — TS-A3-05 to TS-A3-07'),
        ('11', 'Admin Platform — Assessment Config', 'Resend invite (PENDING status only) (A-3.4)', 'Epic A-E3 — TS-A3-08 to TS-A3-09'),
        ('12', 'Admin Platform — Question Mgmt', 'AI question generation with streaming (A-2.1)', 'Epic A-E2 — TS-A2-01 to TS-A2-02'),
        ('13', 'Admin Platform — Question Mgmt', 'Review/edit AI questions; Save Draft / Publish / Regenerate (A-2.2)', 'Epic A-E2 — TS-A2-03 to TS-A2-06'),
        ('14', 'Admin Platform — Question Mgmt', 'Manual question creation with test cases (A-2.3)', 'Epic A-E2 — TS-A2-07 to TS-A2-08'),
        ('15', 'Admin Platform — Question Mgmt', 'Question bank CRUD; versioning; soft-delete guards (A-2.4)', 'Epic A-E2 — TS-A2-09 to TS-A2-12'),
        ('16', 'Admin Platform — Email', 'Invite email content and HMAC link format (A-4.1)', 'Epic A-E4 — TS-A4-01 to TS-A4-02'),
        ('17', 'Admin Platform — Email', 'Failed delivery retry (3x, 1s/4s/16s backoff) + manual resend (A-4.2)', 'Epic A-E4 — TS-A4-03 to TS-A4-04'),
        ('18', 'Admin Platform — Statistics', 'Per-assessment analytics: sortable table, CSV export, histogram (A-5.1)', 'Epic A-E5 — TS-A5-01 to TS-A5-02'),
        ('19', 'Admin Platform — Statistics', 'User-specific statistics across assessments (A-5.2)', 'Epic A-E5 — TS-A5-03'),
        ('20', 'Integration Contract', 'GET /assessments/{id}/config — payload validation', 'INT — TS-INT-01 to TS-INT-02'),
        ('21', 'Integration Contract', 'GET /questions/{id} and GET /assessments/{id}/testcases', 'INT — TS-INT-03 to TS-INT-04'),
        ('22', 'Integration Contract', 'GET /results/assessment/{id} and /results/user/{id} (ADMIN JWT)', 'INT — TS-INT-05 to TS-INT-07'),
        ('23', 'Integration Contract', 'WebSocket RUN/OUTPUT/DONE message contract + cross-session guard', 'INT — TS-INT-08 to TS-INT-09'),
    ]
    for row_data in in_scope:
        row = in_scope_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    add_heading(doc, '2.2 Components and Functions NOT to be Tested', 2)
    out_scope_tbl = doc.add_table(rows=1, cols=3)
    out_scope_tbl.style = 'Table Grid'
    for c, t in zip(out_scope_tbl.rows[0].cells, ['#', 'Component', 'Reason']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    out_scope = [
        ('1', 'Team 2 Candidate Platform internal features (B-E1–B-E5)', 'Owned and tested by Team 2'),
        ('2', 'Execution Engine container orchestration internals', 'Out of QA scope; treated as black box'),
        ('3', 'SendGrid SMTP infrastructure', 'Third-party; mock/stub in QA environment'),
        ('4', 'Spring AI / LLM provider infrastructure', 'Third-party; tested via mocked responses for unit; integration tested with live endpoint'),
        ('5', 'Redis internal cluster operations', 'Infrastructure-level; not in functional test scope'),
    ]
    for row_data in out_scope:
        row = out_scope_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    add_heading(doc, '2.3 Third-Party Components', 2)
    tp_tbl = doc.add_table(rows=1, cols=3)
    tp_tbl.style = 'Table Grid'
    for c, t in zip(tp_tbl.rows[0].cells, ['#', 'Component', 'Role / Comment']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    third_party = [
        ('1', 'SendGrid SMTP', 'Email delivery for invite and password reset emails. Mocked in QA.'),
        ('2', 'Spring AI / LLM provider', 'Powers AI question generation. Requires configurable mock/stub for deterministic testing.'),
        ('3', 'Redis', 'Session caching and token revocation. Required live in QA environment.'),
        ('4', 'Monaco Editor (Microsoft)', 'Code editor component in Admin and Candidate views. Tested via UI scenarios.'),
    ]
    for row_data in third_party:
        row = tp_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ── 3. Quality and Acceptance Criteria ───────────────────────────────────
    add_heading(doc, '3. Quality and Acceptance Criteria', 1)
    qac = [
        'The product must work according to the requirements and functional specifications in the CodEval epics.',
        'No Critical or Major bugs open at the time of release.',
        'All happy-path test scenarios (TS-E1-*, TS-A*-*, TS-INT-*) pass.',
        'All integration contract API tests pass with correct request/response payloads.',
        'RBAC enforcement confirmed on all protected endpoints (isAdmin / isUser middleware).',
        'HMAC-signed token validation verified for both invite links and password reset links.',
    ]
    for item in qac:
        add_bullet(doc, item)

    # ── 4. Critical Success Factors ───────────────────────────────────────────
    add_heading(doc, '4. Critical Success Factors', 1)
    csf = [
        'Complete all Team 1 feature testing within the sprint schedule.',
        'Integration contract API tests pass before Team 2 candidate-facing features are enabled in QA.',
        'Application must not have known Critical or Major bugs at Final Release.',
        'Authentication flows (login, signup, password reset) must be stable before any other module is tested.',
        'AI question generation failure path (retry + Try Again) must be verified with mocked LLM failures.',
    ]
    for item in csf:
        add_bullet(doc, item)

    # ── 5. Risk Assessment ────────────────────────────────────────────────────
    add_heading(doc, '5. Risk Assessment', 1)
    risk_tbl = doc.add_table(rows=1, cols=6)
    risk_tbl.style = 'Table Grid'
    for c, t in zip(risk_tbl.rows[0].cells,
                    ['#', 'Risk', 'Probability %', 'Impact', 'Preventive Actions', 'Contingency']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    risks = [
        ('1', 'Integration contract API schema drift between Team 1 and Team 2', '40', 'High',
         'Contract tests automated and run on every build', 'Freeze contract early; flag discrepancy immediately to both leads'),
        ('2', 'LLM provider unavailability affecting AI generation tests', '30', 'Medium',
         'Use mock/stub LLM in QA for deterministic testing', 'Tests that require live LLM gated to integration phase only'),
        ('3', 'HMAC token clock-sync issue causing false expiry failures', '20', 'Medium',
         'Use server-side time consistently; document clock-sync requirement', 'Add configurable buffer to expiry checks in QA'),
        ('4', 'Monaco editor compatibility issues in unsupported browsers', '25', 'Medium',
         'Define browser compatibility matrix up front', 'Flag as Known Limitation; exclude from release blocker'),
        ('5', 'SendGrid SMTP failures in QA environment', '35', 'Low',
         'Mock SendGrid in QA; use real relay only in staging', 'Use mock assertions for email content; live send in staging only'),
    ]
    for row_data in risks:
        row = risk_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ── 6. Resources ──────────────────────────────────────────────────────────
    add_heading(doc, '6. Resources', 1)
    add_heading(doc, '6.1 Key Project Resources', 2)
    res_tbl = doc.add_table(rows=1, cols=3)
    res_tbl.style = 'Table Grid'
    for c, t in zip(res_tbl.rows[0].cells, ['#', 'Project Role', 'Name / Email / Location']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row_data in [
        ('1', 'Project Manager', '<TBD>'),
        ('2', 'Team 1 Lead Developer', '<TBD>'),
        ('3', 'Team 2 Lead Developer', '<TBD>'),
        ('4', 'QA Test Lead', '<TBD>'),
    ]:
        row = res_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    add_heading(doc, '6.2 Test Team', 2)
    team_tbl = doc.add_table(rows=1, cols=4)
    team_tbl.style = 'Table Grid'
    for c, t in zip(team_tbl.rows[0].cells, ['#', 'Role', 'Name', 'Responsibilities']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row_data in [
        ('1', 'QA Lead', '<TBD>', 'Test strategy, plan, reporting, integration test co-ordination'),
        ('2', 'QA Engineer (Team 1)', '<TBD>', 'Auth, Admin Platform feature testing'),
        ('3', 'QA Engineer (Integration)', '<TBD>', 'API contract and integration tests across teams'),
    ]:
        row = team_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    add_heading(doc, '6.3 Test Environment', 2)
    env_tbl = doc.add_table(rows=1, cols=4)
    env_tbl.style = 'Table Grid'
    for c, t in zip(env_tbl.rows[0].cells, ['#', 'Environment', 'Purpose', 'Configuration']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row_data in [
        ('1', 'DEV (local Docker Compose)', 'Developer self-testing; unit + component tests', 'Java 21, Spring Boot, Redis, H2/PostgreSQL'),
        ('2', 'QA Staging', 'All functional, API, and integration tests', 'Both Team 1 and Team 2 services deployed; PostgreSQL; Redis; mocked SendGrid'),
        ('3', 'UAT', 'User acceptance (if applicable)', 'Production-like environment; live SendGrid'),
    ]:
        row = env_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    add_heading(doc, '6.3.1 Test Tools', 3)
    tool_tbl = doc.add_table(rows=1, cols=2)
    tool_tbl.style = 'Table Grid'
    for c, t in zip(tool_tbl.rows[0].cells, ['Tool', 'Purpose']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row_data in [
        ('JIRA', 'Bug tracking and test execution tracking'),
        ('QA Space (Jira plugin)', 'Test case and documentation management'),
        ('Postman / REST Assured', 'API and integration contract testing'),
        ('Selenium / Playwright', 'UI smoke and regression automation'),
        ('WireMock / MockServer', 'Mocking SendGrid and LLM provider in QA'),
    ]:
        row = tool_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ── 7. Test Documentation and Deliverables ────────────────────────────────
    add_heading(doc, '7. Test Documentation and Deliverables', 1)
    deliv_tbl = doc.add_table(rows=1, cols=4)
    deliv_tbl.style = 'Table Grid'
    for c, t in zip(deliv_tbl.rows[0].cells, ['#', 'Document', 'Responsible', 'Delivery']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row_data in [
        ('1', 'CodEval Test Strategy (this document)', 'QA Lead', 'Before testing starts'),
        ('2', 'CodEval Test Plan (this document)', 'QA Lead', 'Before testing starts'),
        ('3', 'QA Output JSON (qa_output.json)', 'QA Engineer', 'Before Phase 2'),
        ('4', 'Test Case Execution Reports', 'QA Engineers', 'After each sprint'),
        ('5', 'Bug Reports (JIRA)', 'QA Engineers', 'On each defect found'),
        ('6', 'Weekly Test Status Reports', 'QA Lead', 'Weekly during testing'),
        ('7', 'Final Test Summary Report', 'QA Lead', 'Before release'),
    ]:
        row = deliv_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ── 8. Test Strategy ──────────────────────────────────────────────────────
    add_heading(doc, '8. Test Strategy', 1)
    add_para(doc, (
        'The CodEval applications will be tested using a black-box approach based on requirements '
        'and acceptance criteria without knowledge of internal implementation. Where API contracts '
        'are tested, grey-box testing is applied using the agreed integration contract schemas.'
    ))

    add_heading(doc, '8.1 Entry Criteria', 2)
    add_para(doc, 'Testing may be suspended if any of the following occurs:')
    for item in [
        'The new build cannot be installed successfully.',
        'A critical fault prevents testing of the feature under test.',
        'The application does not contain the specified change(s).',
        'New claimed functionality does not work or works improperly.',
        'A severe problem prevents testing from continuing.',
        'Development has not corrected previously identified suspension issues.',
        'A new build version is released during active testing.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '8.2 Test Methods', 2)
    for item in [
        'Manual functional testing — primary method for all feature and UI tests.',
        'Automated API testing — Postman collections / REST Assured for all integration contract endpoints.',
        'Automated smoke testing — runs on every QA build deployment.',
        'Exploratory testing — time-boxed sessions targeting auth flows and admin lifecycle scenarios.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '8.3 Test Types', 2)
    for item in [
        'Functional Testing — all Team 1 user stories',
        'API / Integration Testing — INT epics (cross-team contracts)',
        'Smoke Testing — critical path on every build',
        'Regression Testing — after each bug fix batch',
        'UI Testing — admin platform pages',
        'Exploratory / Ad-hoc Testing — auth, AI generation, lifecycle edge cases',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '8.4 Test Levels', 2)
    add_heading(doc, '8.4.1 Smoke Test', 3)
    add_para(doc, (
        'Smoke tests verify the most critical build functions: login with valid EPAM credentials '
        '(TS-E1-01), admin dashboard load, user list display (TS-A1-01), and assessment creation '
        '(TS-A3-01). If smoke tests fail, testing is suspended and the build is rejected.'
    ))
    add_heading(doc, '8.4.2 Critical Path Test', 3)
    add_para(doc, (
        'Critical path tests cover all happy-path scenarios across all 23 features in qa_output.json. '
        'These are run after smoke tests pass. Focus areas: auth flows, assessment lifecycle '
        '(create → invite → ACTIVE), invite email delivery, and integration contract APIs.'
    ))
    add_heading(doc, '8.4.3 Extended Test', 3)
    add_para(doc, (
        'Extended tests cover all explicitly-stated edge cases and negative scenarios in qa_output.json, '
        'including: non-@epam.com domain blocking, single-use HMAC link rejection, soft-delete guards, '
        'AI retry logic, WebSocket cross-session rejection, and RBAC enforcement on all endpoints.'
    ))

    add_heading(doc, '8.5 Bug and Documentation Tracking', 2)
    add_para(doc, (
        'All defects are logged in JIRA with severity, steps to reproduce, expected vs actual results, '
        'and environment details. Test documentation is maintained in QA Space.'
    ))
    add_heading(doc, '8.5.1 Bug Severity Definitions', 3)
    for line in [
        'Critical — Application, component, or module crash or inaccessible.',
        'Major — Data corruption/loss, major functionality broken, no workaround.',
        'Medium — Problem with a workaround; secondary features impaired.',
        'Minor — Cosmetic flaw.',
    ]:
        add_bullet(doc, line)

    # ── 9. Testing Schedule ───────────────────────────────────────────────────
    add_heading(doc, '9. Testing Schedule', 1)
    sched_tbl = doc.add_table(rows=1, cols=5)
    sched_tbl.style = 'Table Grid'
    for c, t in zip(sched_tbl.rows[0].cells,
                    ['#', 'Activity', 'Start Date', 'End Date', 'Responsible']):
        c.text = t
        for run in c.paragraphs[0].runs:
            run.bold = True
    schedule = [
        ('1', 'Test Plan & Strategy creation', '<TBD>', '<TBD>', 'QA Lead'),
        ('2', 'Test case creation (qa_output.json)', '<TBD>', '<TBD>', 'QA Engineers'),
        ('3', 'QA environment setup and smoke test', '<TBD>', '<TBD>', 'QA Lead'),
        ('4', 'Phase 2: Feature testing — Auth + Admin Platform', '<TBD>', '<TBD>', 'QA Engineer (Team 1)'),
        ('5', 'Phase 3: Integration contract API testing', '<TBD>', '<TBD>', 'QA Engineer (Integration)'),
        ('6', 'Phase 4: Regression testing', '<TBD>', '<TBD>', 'QA Engineers'),
        ('7', 'UAT (if applicable)', '<TBD>', '<TBD>', 'QA Lead + Stakeholders'),
        ('8', 'Final Test Summary Report', '<TBD>', '<TBD>', 'QA Lead'),
    ]
    for row_data in schedule:
        row = sched_tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    doc.save(dst)
    print(f'Test Plan saved: {dst}')


if __name__ == '__main__':
    build_test_strategy()
    build_test_plan()
    print('All documents generated successfully.')
